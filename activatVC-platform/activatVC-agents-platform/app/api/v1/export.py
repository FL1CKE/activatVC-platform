from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.services.run_service import RunService

router = APIRouter()


@router.get("/run/{run_id}/md")
async def export_run_markdown(run_id: int, db: AsyncSession = Depends(get_db)):
    """Скачать сводный отчёт всей команды агентов (full run) как Markdown."""
    run = await RunService(db).get_run_with_tasks(run_id)
    if not run:
        raise HTTPException(404, f"Run {run_id} not found")

    parts = [f"# Startup Analysis Report\n\n**Startup:** {run.startup_name or run.application_id}\n**Run ID:** {run.id}\n**Status:** {run.status}\n"]
    for task in sorted(run.tasks or [], key=lambda t: t.agent.role if t.agent else ""):
        role = task.agent.role if task.agent else "Agent"
        name = task.agent.name if task.agent else ""
        parts.append(f"\n---\n\n## {role} — {name}\n\n{task.report_content or '*No report*'}\n")

    content = "\n".join(parts)
    filename = f"run_{run_id}_full_report.md"
    return Response(
        content=content.encode("utf-8"),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{task_id}/md")
async def export_markdown(task_id: int, db: AsyncSession = Depends(get_db)):
    """Скачать отчёт агента как Markdown файл."""
    task = await RunService(db).get_task(task_id)
    if not task:
        raise HTTPException(404, f"Task {task_id} not found")
    if not task.report_content:
        raise HTTPException(404, "Report not ready yet")

    role = task.agent.role if task.agent else "agent"
    filename = f"{role}_report_{task_id}.md"

    return Response(
        content=task.report_content.encode("utf-8"),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{task_id}/docx")
async def export_docx(task_id: int, db: AsyncSession = Depends(get_db)):
    """Скачать отчёт как DOCX файл."""
    task = await RunService(db).get_task(task_id)
    if not task:
        raise HTTPException(404, f"Task {task_id} not found")
    if not task.report_content:
        raise HTTPException(404, "Report not ready yet")

    try:
        content = _md_to_docx(task.report_content)
    except Exception as e:
        raise HTTPException(500, f"DOCX generation failed: {e}")

    role = task.agent.role if task.agent else "agent"
    filename = f"{role}_report_{task_id}.docx"

    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{task_id}/pdf")
async def export_pdf(task_id: int, db: AsyncSession = Depends(get_db)):
    """Скачать отчёт как PDF файл."""
    task = await RunService(db).get_task(task_id)
    if not task:
        raise HTTPException(404, f"Task {task_id} not found")
    if not task.report_content:
        raise HTTPException(404, "Report not ready yet")

    try:
        content = _md_to_pdf(task.report_content)
    except Exception as e:
        raise HTTPException(500, f"PDF generation failed: {e}")

    role = task.agent.role if task.agent else "agent"
    filename = f"{role}_report_{task_id}.pdf"

    return Response(
        content=content,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Converters ────────────────────────────────────────────────────────────────

def _md_to_docx(markdown_text: str) -> bytes:
    """
    Конвертация Markdown → DOCX.
    Парсим базовые элементы: заголовки, параграфы, bold, bullet lists.
    """
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Стили заголовков
    for level, name in [(1, "Heading 1"), (2, "Heading 2"), (3, "Heading 3")]:
        style = doc.styles[name]
        style.font.size = Pt(18 - (level - 1) * 2)

    for line in markdown_text.split("\n"):
        line_stripped = line.strip()
        if not line_stripped:
            doc.add_paragraph()
            continue

        if line_stripped.startswith("### "):
            doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith("## "):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith("# "):
            doc.add_heading(line_stripped[2:], level=1)
        elif line_stripped.startswith(("- ", "* ", "• ")):
            doc.add_paragraph(line_stripped[2:], style="List Bullet")
        elif line_stripped.startswith(tuple(f"{i}." for i in range(1, 20))):
            text = line_stripped.split(".", 1)[1].strip()
            doc.add_paragraph(text, style="List Number")
        else:
            # Обрабатываем **bold** в параграфе
            para = doc.add_paragraph()
            _add_runs_with_bold(para, line_stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs_with_bold(para, text: str):
    """Разбивает текст на runs с bold форматированием для **text**."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (i % 2 == 1)  # нечётные части — внутри **...**


def _md_to_pdf(markdown_text: str) -> bytes:
    """
    Конвертация Markdown → PDF через HTML промежуточный формат.
    markdown → HTML → PDF (xhtml2pdf). Поддержка кириллицы через Arial/DejaVu.
    """
    import io
    import os
    import markdown as md_lib
    from xhtml2pdf import pisa
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Регистрируем шрифт с поддержкой кириллицы
    _FONT_CANDIDATES = [
        ("ArialCyr", r"C:\Windows\Fonts\arial.ttf"),
        ("ArialCyr", r"C:\Windows\Fonts\Arial.ttf"),
        ("ArialCyr", "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
        ("ArialCyr", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        ("ArialCyr", "/usr/share/fonts/dejavu/DejaVuSans.ttf"),
    ]
    font_name = "Helvetica"
    for _name, _path in _FONT_CANDIDATES:
        if os.path.exists(_path):
            try:
                pdfmetrics.registerFont(TTFont(_name, _path))
                font_name = _name
                break
            except Exception:
                pass

    html_body = md_lib.markdown(
        markdown_text,
        extensions=["tables", "fenced_code", "nl2br"],
    )

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @page {{ margin: 20mm 18mm 22mm 18mm; }}
  body {{ font-family: {font_name}; font-size: 11px; line-height: 1.65;
         color: #1e293b; background: #fff; }}
  /* Header band */
  .brand-header {{ background: #0d2817; color: #22c55e; padding: 10px 16px;
                   margin: -20px -18px 20px -18px; border-bottom: 3px solid #22c55e; }}
  .brand-header .brand {{ font-size: 16px; font-weight: bold; color: #22c55e; }}
  .brand-header .subtitle {{ font-size: 9px; color: #a7f3d0; margin-top: 2px; }}
  /* Headings */
  h1 {{ font-size: 18px; color: #0d2817; border-bottom: 2px solid #22c55e;
        padding-bottom: 6px; margin: 20px 0 10px; }}
  h2 {{ font-size: 14px; color: #1a4a2e; border-left: 3px solid #22c55e;
        padding-left: 8px; margin: 18px 0 8px; }}
  h3 {{ font-size: 12px; color: #166534; margin: 14px 0 6px; }}
  h4 {{ font-size: 11px; color: #475569; margin: 10px 0 4px; }}
  /* Tables */
  table {{ border-collapse: collapse; width: 100%; margin: 14px 0; font-size: 10px; }}
  thead tr {{ background: #1a4a2e; color: #ffffff; }}
  th {{ padding: 7px 10px; text-align: left; font-weight: bold; border: 1px solid #166534; }}
  td {{ border: 1px solid #d1fae5; padding: 6px 10px; vertical-align: top; }}
  tr:nth-child(even) td {{ background: #f0fdf4; }}
  tr:nth-child(odd) td {{ background: #ffffff; }}
  /* Lists */
  ul {{ margin: 6px 0; padding-left: 20px; }}
  li {{ margin: 3px 0; line-height: 1.5; }}
  ul li::marker {{ color: #22c55e; }}
  /* Code */
  code {{ background: #f1f5f9; padding: 1px 5px; border-radius: 3px;
          font-size: 10px; color: #475569; }}
  pre {{ background: #f1f5f9; padding: 10px 14px; border-radius: 4px;
         border-left: 3px solid #22c55e; overflow-x: auto; font-size: 9.5px; }}
  /* Special text */
  strong {{ color: #0d2817; font-weight: bold; }}
  em {{ color: #475569; }}
  /* Score/verdict badges via bold patterns */
  p {{ margin: 6px 0; }}
  blockquote {{ border-left: 3px solid #22c55e; padding: 6px 12px;
                background: #f0fdf4; color: #166534; margin: 10px 0; }}
  /* Footer */
  .page-footer {{ position: fixed; bottom: -15mm; left: 0; right: 0; height: 12mm;
                  font-size: 8px; color: #94a3b8; border-top: 1px solid #e2e8f0;
                  padding-top: 3px; text-align: center; }}
  @page {{ @bottom-center {{ content: "Activat VC · Конфиденциально · Страница " counter(page) " из " counter(pages); font-size: 8px; color: #94a3b8; }} }}
</style>
</head>
<body>
<div class="brand-header">
  <div class="brand">ACTIVAT VC</div>
  <div class="subtitle">Due Diligence Platform · Автоматизированный анализ стартапов</div>
</div>
{html_body}
</body>
</html>"""

    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(html, dest=buffer, encoding="utf-8")
    if pisa_status.err:
        raise RuntimeError(f"xhtml2pdf conversion failed with {pisa_status.err} errors")
    return buffer.getvalue()
