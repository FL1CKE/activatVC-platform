"""
Report Export Service.

Конвертирует markdown-отчёт агента в нужный формат.
Все агенты пишут отчёты в Markdown — это единый формат.
Отсюда конвертируем в DOCX и PDF по запросу.
"""
import io
import logging
from enum import Enum

logger = logging.getLogger(__name__)


class ExportFormat(str, Enum):
    MARKDOWN = "md"
    DOCX = "docx"
    PDF = "pdf"


def export_report(content: str, fmt: ExportFormat, title: str = "Agent Report") -> bytes:
    """
    Конвертирует markdown строку в нужный формат.
    Возвращает bytes готового файла.
    """
    if fmt == ExportFormat.MARKDOWN:
        return content.encode("utf-8")
    elif fmt == ExportFormat.DOCX:
        return _to_docx(content, title)
    elif fmt == ExportFormat.PDF:
        return _to_pdf(content, title)
    else:
        raise ValueError(f"Unknown export format: {fmt}")


def get_content_type(fmt: ExportFormat) -> str:
    return {
        ExportFormat.MARKDOWN: "text/markdown; charset=utf-8",
        ExportFormat.DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ExportFormat.PDF: "application/pdf",
    }[fmt]


def get_file_extension(fmt: ExportFormat) -> str:
    return fmt.value


# ─── Converters ────────────────────────────────────────────────────────────────

def _to_docx(markdown_text: str, title: str) -> bytes:
    """Markdown → DOCX через python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Стили
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Парсим markdown построчно — простой конвертер
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("**") and line.endswith("**") and len(line) > 4:
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        elif line.strip() == "---":
            doc.add_paragraph("─" * 60)
        elif line.strip():
            # Обычный параграф — обрабатываем inline **bold**
            p = doc.add_paragraph()
            _add_inline_formatted(p, line)
        else:
            # Пустая строка
            doc.add_paragraph("")

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_inline_formatted(paragraph, text: str):
    """Обрабатывает **bold** внутри строки."""
    import re
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _to_pdf(markdown_text: str, title: str) -> bytes:
    """
    Markdown → PDF через xhtml2pdf (HTML → PDF).
    Чистый Python, работает на Windows/Linux/Mac без нативных зависимостей.
    """
    try:
        import io
        import markdown as md_lib
        from xhtml2pdf import pisa

        # Конвертируем markdown в HTML
        html_content = md_lib.markdown(
            markdown_text,
            extensions=["tables", "fenced_code", "nl2br"],
        )

        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{title}</title>
            <style>
                body {{
                    font-family: Helvetica, Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    margin: 2cm;
                    color: #1a1a1a;
                }}
                h1 {{ font-size: 20pt; color: #1a1a2e; border-bottom: 2px solid #1a1a2e; padding-bottom: 8px; }}
                h2 {{ font-size: 15pt; color: #16213e; margin-top: 24px; }}
                h3 {{ font-size: 12pt; color: #0f3460; }}
                table {{ border-collapse: collapse; width: 100%; margin: 12px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background: #f0f0f0; font-weight: bold; }}
                code {{ background: #f5f5f5; padding: 2px 4px; font-size: 10pt; }}
                hr {{ border: none; border-top: 1px solid #ccc; margin: 16px 0; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(full_html, dest=buffer, encoding="utf-8")
        if pisa_status.err:
            raise RuntimeError(f"xhtml2pdf conversion failed with {pisa_status.err} errors")
        return buffer.getvalue()

    except ImportError:
        logger.warning("xhtml2pdf not available, falling back to basic PDF via reportlab")
        return _to_pdf_reportlab(markdown_text, title)


def _to_pdf_reportlab(text: str, title: str) -> bytes:
    """Фолбэк PDF через reportlab если WeasyPrint недоступен."""
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    for line in text.split("\n"):
        if line.startswith("# "):
            story.append(Paragraph(line[2:], styles["h1"]))
        elif line.startswith("## "):
            story.append(Paragraph(line[3:], styles["h2"]))
        elif line.strip():
            story.append(Paragraph(line.replace("**", "<b>", 1).replace("**", "</b>", 1), styles["Normal"]))
        story.append(Spacer(1, 4))

    doc.build(story)
    return buffer.getvalue()
